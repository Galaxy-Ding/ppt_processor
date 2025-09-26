from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path
import comtypes.client
import time
from PIL import Image
from typing import Dict, Any, Tuple
from utils.logger import LoggerFactory
# 获取当前文件的绝对路径的根目录
current_file_path = Path(__file__).resolve()
main_dir = current_file_path.parent.parent  # 项目根目录（ppt_processor/）


class ImageExporter:
    """图片导出处理类"""
    def __init__(self, pptx_path: str):
        """
        初始化导出器
        
        Args:
            pptx_path: PowerPoint文件路径
        """
        self.logger = LoggerFactory.create_logger("ImageExporter")
        self.pptx_path = pptx_path
        self.temp_dir = os.path.join(main_dir, ".temp")
        self._ensure_temp_dir()
        self.logger.debug(f"temp_dir: {self.temp_dir}")
        self.logger.debug(f"pptx_path: {self.pptx_path}")
        
        # 初始化PowerPoint类型库
        try:
            from comtypes.gen import PowerPoint
        except ImportError:
            powerpoint = comtypes.client.CreateObject("PowerPoint.Application")
            powerpoint.Quit()
            from comtypes.gen import PowerPoint
    
    def _export_slide_region_as_image(self, slide_index: int, output_path: str, 
                                    region_cm: Tuple[float, float, float, float]) -> bool:
        """
        将PPT中指定页面的特定区域导出为图片
        
        Args:
            slide_index: 要导出的幻灯片索引(从1开始)
            output_path: 图片保存路径
            region_cm: 区域坐标(left, top, width, height)，单位为厘米
        """
        powerpoint = None
        try:
            powerpoint = comtypes.client.CreateObject("PowerPoint.Application")
            presentation = powerpoint.Presentations.Open(
                self.pptx_path,
                WithWindow=False
            )
            
            # 导出完整幻灯片
            temp_output = output_path
            slide = presentation.Slides[slide_index]
            slide.Export(temp_output, "PNG")
            presentation.Close()
            
            # 裁剪指定区域
            with Image.open(temp_output) as img:
                # 将厘米转换为像素 (96 DPI)
                left_cm, top_cm, width_cm, height_cm = region_cm
                dpi = 96
                pixels_per_cm = dpi / 2.54
                
                left_px = int(left_cm * pixels_per_cm)
                top_px = int(top_cm * pixels_per_cm)
                width_px = int(width_cm * pixels_per_cm)
                height_px = int(height_cm * pixels_per_cm)
                
                # 裁剪并保存
                cropped_img = img.crop((
                    left_px, top_px, 
                    left_px + width_px, 
                    top_px + height_px
                ))
                cropped_img.save(output_path)
            
            return True
            
        except Exception as e:
            print(f"Error exporting slide region: {e}")
            return False
        finally:
            if powerpoint:
                try:
                    powerpoint.Quit()
                    time.sleep(1)
                except:
                    pass

    def _ensure_temp_dir(self):
        if not os.path.exists(self.temp_dir):
            os.makedirs(self.temp_dir)
    
    def _clean_temp_files(self):
        # if self.logger.log_level < 20:
        #     if self.logger:
        #         self.logger.debug("DEBUG等级下，跳过临时文件删除。")
        #     return
        for file in os.listdir(self.temp_dir):
            file_path = os.path.join(self.temp_dir, file)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
            except Exception as e:
                print(f"Error deleting {file_path}: {e}")

    def export_images(self, extracted_data: Dict[str, Any]) -> Dict[str, str]:
        image_paths = {}
        # 遍历提取的数据
        for field_name, data in extracted_data.items():
            if isinstance(data, dict) and "box" in data and "page_number" in data:
                temp_file = os.path.join(self.temp_dir, f"{field_name}.png")
                success = self._export_slide_region_as_image(
                    slide_index=data["page_number"],
                    output_path=temp_file,
                    region_cm=data["box"]
                )
                if success:
                    image_paths[field_name] = temp_file
                else:
                    print(f"Failed to export image for {field_name}")
        return image_paths

class DocxProcessor:
    """Word文档处理类"""
    def __init__(self, docx_path: str, output_path: str):
        self.docx_path = docx_path
        self.output_path = output_path
        self.doc = None
        self.logger = LoggerFactory.create_logger("DocxProcessor")
        
    def process_content(self, replacements: Dict[str, Any]) -> bool:
        try:
            self.doc = Document(self.docx_path)
            text_replacements = {}
            image_mappings = {}
            
            for key, value in replacements.items():
                if isinstance(value, dict) and 'path' in value:
                    image_mappings[key] = value
                else:
                    text_replacements[key] = str(value)
            
            if text_replacements:
                self._process_all_text(text_replacements)
            
            if image_mappings:
                self._insert_images(image_mappings)
            
            self.doc.save(self.output_path)
            return True
            
        except Exception as e:
            print(f"Error processing document: {e}")
            return False

    def _process_all_text(self, replacements: Dict[str, str]) -> None:
        for section in self.doc.sections:
            for para in section.header.paragraphs:
                self._replace_keywords_in_paragraph(para, replacements)
            for para in section.footer.paragraphs:
                self._replace_keywords_in_paragraph(para, replacements)
            for table in section.header.tables:
                self._replace_keywords_in_table(table, replacements)
            for table in section.footer.tables:
                self._replace_keywords_in_table(table, replacements)

        for para in self.doc.paragraphs:
            self._replace_keywords_in_paragraph(para, replacements)

        for table in self.doc.tables:
            self._replace_keywords_in_table(table, replacements)

    def _replace_keywords_in_paragraph(self, para, replacements: Dict[str, str]) -> None:
        for run in para.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(value))

    def _replace_keywords_in_table(self, table, replacements: Dict[str, str]) -> None:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    self._replace_keywords_in_paragraph(para, replacements)
                for nested_table in cell.tables:
                    self._replace_keywords_in_table(nested_table, replacements)
    
    def _insert_images(self, image_mappings: Dict[str, Dict]) -> None:
        found_markers = set()
        
        for paragraph in self.doc.paragraphs:
            for marker, img_config in image_mappings.items():
                if marker in paragraph.text:
                    paragraph.clear()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    img_path = img_config['path']
                    img_width = img_config['width']
                    
                    try:
                        run = paragraph.add_run()
                        with open(img_path, 'rb') as f:
                            run.add_picture(f, width=Cm(img_width))
                        found_markers.add(marker)
                    except Exception as e:
                        self.logger.error(f"添加图片失败 {marker}: {e}")
                    break
        
        not_found = set(image_mappings.keys()) - found_markers
        if not_found:
            self.logger.warning(f"以下标记未找到: {', '.join(not_found)}")

class ExporterA:
    """发包规范导出器"""
    def __init__(self, pptx_path: str, docx_template_path: str, output_path: str):
        self.logger = LoggerFactory.create_logger("Exporter发包规范")
        self.logger.info("初始化导出器")
        self.output_path = output_path
        self.image_exporter = ImageExporter(pptx_path)
        self.docx_processor = DocxProcessor(docx_template_path, self.output_path)
        self.logger.debug("pptx_path: %s", pptx_path)
        self.logger.debug("output_path: %s", output_path)
        self.logger.debug("docx_template_path: %s", docx_template_path)

    def process(self, result_data: Dict[str, Any]) -> bool:
        try:
            self.logger.info("开始处理文档")
            
            # 1. 导出图片到临时目录
            self.logger.debug("导出图片到临时目录")
            image_paths = self.image_exporter.export_images(result_data)
            
            # 2. 将图片和文本内容插入到文档
            self.logger.debug("处理替换内容")
            replacements = {}
            # 添加非图片字段
            replacements.update({
                k: v for k, v in result_data.items() 
                if not (isinstance(v, dict) and 'box' in v and 'page_number' in v)
            })
            # 添加图片字段，使用shape中的宽度
            replacements.update({
                k: {
                    'path': path,
                    'width': result_data[k]['box'][2]  # 使用box中的宽度
                }
                for k, path in image_paths.items()
            })
            
            self.docx_processor.process_content(replacements)
            
            self.logger.info(f"文档生成成功: {self.output_path}")
            # 3. 清理临时文件
            self.image_exporter._clean_temp_files()
            
            return True
            
        except Exception as e:
            self.logger.error(f"处理文档时出错: {e}", exc_info=True)
            self.image_exporter._clean_temp_files()
            return False

if __name__ == "__main__":
    # 使用示例
    pptx_path = r"D:\pythonf\ppt_processor\examples\templates\26xdemo2.pptx"
    docx_path = r"..\examples/templates/26xdemo1.docx"
    output_path = r"..\output.docx"
    
    # 创建导出器实例
    exporter = ExporterA(pptx_path, docx_path, output_path)
    
    # 测试数据
    extracted_data = {
        "img_dev_frontlooking": {"box": (15.16, 9.38, 11.73, 8.28), "page_number": 7},
        "img_dev_occupancy": {"box": (15.16, 9.38, 11.73, 8.28), "page_number": 7},
        "img_dev_overlooking": {"box": (15.16, 9.38, 11.73, 8.28), "page_number": 7},
        "img_dev_craftsmanship": {"box": (15.16, 9.38, 11.73, 8.28), "page_number": 7}
    }
    
    success = exporter.process(extracted_data)
    print(f"文档生成{'成功' if success else '失败'}")